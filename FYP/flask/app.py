from collections import Counter
import spacy
import weasyprint
import gensim.corpora as corpora
import gensim
from ibm_cloud_sdk_core.authenticators import IAMAuthenticator
from ibm_watson import ToneAnalyzerV3
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
import base64
import matplotlib.pyplot as plt
from textblob import TextBlob
from heapq import nlargest
from sklearn.feature_extraction.stop_words import ENGLISH_STOP_WORDS as stop_words
from string import punctuation
from itsdangerous import URLSafeTimedSerializer
from flask_mail import Mail, Message
from authlib.integrations.flask_client import OAuth
from flask_dance.contrib.github import make_github_blueprint, github
from werkzeug.utils import secure_filename
from mobi import mobi_uncompress
import nltk
from flask import (
    Flask,
    render_template,
    request,
    session,
    logging,
    url_for,
    redirect,
    flash,
)
import requests
from sqlalchemy import create_engine
from sqlalchemy.orm import scoped_session, sessionmaker
from passlib.hash import sha256_crypt
from flask import send_from_directory
import PyPDF2
import pytesseract
import io
from PIL import Image
import docx
from bs4 import BeautifulSoup
from urllib.request import urlopen
from ebooklib import epub
import ebooklib
import re
import os
from flask import send_file

os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"

engine = create_engine("")
db = scoped_session(sessionmaker(bind=engine))
app = Flask(__name__)
app.config.from_pyfile("config.cfg")
mail = Mail(app)


user_id = 1
page = ""
summ = ""
positive_sentences = ""
negative_sentences = ""
neutral_sentences = ""
positive = ""
negative = ""
neutral = ""
x = []
y = []
f_name = ""
emails = ""
emotions = ""
intensity = ""
pages_length = 1
topic_pages = []
topicss = []

UPLOAD_FOLDER = "/FYP/flask/data/files"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

apikey = ""
url = ""
authenticator = IAMAuthenticator(apikey)
ta = ToneAnalyzerV3(version="2017-09-21", authenticator=authenticator)
ta.set_service_url(url)


github_blueprint = make_github_blueprint(
    client_id="",
    client_secret="",
)
app.register_blueprint(github_blueprint, url_prefix="/github_login")

oauth = OAuth(app)
google = oauth.register(
    name="google",
    client_id=""
    client_secret="",
    access_token_url="https://accounts.google.com/o/oauth2/token",
    access_token_params=None,
    authorize_url="https://accounts.google.com/o/oauth2/auth",
    authorize_params=None,
    api_base_url="https://www.googleapis.com/oauth2/v1/",
    # This is only needed if using openId to fetch user info
    userinfo_endpoint="https://openidconnect.googleapis.com/v1/userinfo",
    client_kwargs={"scope": "openid email profile"},
)


@app.route("/")
def front():
    return render_template("./sidebar-09/entrance.html")


@app.route("/github-login")
def github_login():
    global user_id
    if not github.authorized:
        return redirect(url_for("github.login"))
    else:
        account_info = github.get("/user")
        if account_info.ok:
            account_info_json = account_info.json()
            user = account_info_json["id"]
            user1 = account_info_json["login"]
            usernamedata = db.execute(
                "SELECT user_name FROM users WHERE user_name=:user_name",
                {"user_name": user},
            ).fetchone()
            if usernamedata is None:
                db.execute(
                    "INSERT INTO users(name, user_name, profession, password) VALUES(:name, :user_name, :profession, :password)",
                    {
                        "name": user1,
                        "user_name": user,
                        "profession": "none",
                        "password": "none",
                    },
                )
                db.commit()
                user_id = db.execute(
                    "SELECT id FROM users WHERE user_name=:user_name",
                    {"user_name": user},
                ).fetchone()
                return redirect(url_for("home"))
            else:
                user_id = db.execute(
                    "SELECT id FROM users WHERE user_name=:user_name",
                    {"user_name": user},
                ).fetchone()
                return redirect(url_for("home"))
            return redirect(url_for("home"))
    return "<h1>Request failed!</h1> "


@app.route("/google-login")
def google_login():
    google = oauth.create_client("google")
    redirect_uri = url_for("authorize", _external=True)
    return google.authorize_redirect(redirect_uri)


@app.route("/authorize")
def authorize():
    global user_id
    google = oauth.create_client("google")
    token = google.authorize_access_token()
    resp = google.get("userinfo")
    user_info = resp.json()
    user = user_info["email"]
    user1 = user_info["name"]
    usernamedata = db.execute(
        "SELECT user_name FROM users WHERE user_name=:user_name", {
            "user_name": user}
    ).fetchone()
    if usernamedata is None:
        db.execute(
            "INSERT INTO users(name, user_name, profession, password) VALUES(:name, :user_name, :profession, :password)",
            {
                "name": user1,
                "user_name": user,
                "profession": "none",
                "password": "none",
            },
        )
        db.commit()
        user_id = db.execute(
            "SELECT id FROM users WHERE user_name=:user_name", {
                "user_name": user}
        ).fetchone()
        return redirect(url_for("home"))
    else:
        user_id = db.execute(
            "SELECT id FROM users WHERE user_name=:user_name", {
                "user_name": user}
        ).fetchone()
        return redirect(url_for("home"))
    return redirect(url_for("home"))


@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        name = request.form.get("fullname")
        user = request.form.get("username")
        profession = request.form.get("profession")
        password = request.form.get("password")
        confirm = request.form.get("confirmpassword")
        secure_password = sha256_crypt.encrypt(str(password))

        usernamedata = db.execute(
            "SELECT user_name FROM users WHERE user_name=:user_name",
            {"user_name": user},
        ).fetchone()

        if usernamedata is None:
            if password == confirm:
                db.execute(
                    "INSERT INTO users(name, user_name, profession, password) VALUES(:name, :user_name, :profession, :password)",
                    {
                        "name": name,
                        "user_name": user,
                        "profession": profession,
                        "password": secure_password,
                    },
                )
                db.commit()
                flash("You are registered and can login now!", "success")
                return redirect(url_for("login"))
            else:
                flash("Password does not match!", "danger")
                return render_template("./Authentication/signup.html")
        else:
            flash("user name already taken!", "danger")
            return render_template("./Authentication/signup.html")
    return render_template("./Authentication/signup.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        global user_id
        user_name = request.form.get("username1")
        password = request.form.get("password1")

        usernamedata = db.execute(
            "SELECT user_name FROM users WHERE user_name=:user_name",
            {"user_name": user_name},
        ).fetchone()
        passwordata = db.execute(
            "SELECT password FROM users WHERE user_name=:user_name",
            {"user_name": user_name},
        ).fetchone()
        user_id = db.execute(
            "SELECT id FROM users WHERE user_name=:user_name", {
                "user_name": user_name}
        ).fetchone()
        if usernamedata is None:
            flash("User not Found!", "danger")
            return render_template("./Authentication/login.html")
        else:
            for passwor_data in passwordata:
                if sha256_crypt.verify(password, passwor_data):
                    flash(
                        "Congratulations! You have logged in successfully and can now use Data Science Toolkit.",
                        "success",
                    )
                    return redirect(url_for("home"))
                else:
                    flash("Incorrect password!", "danger")
                    return render_template("./Authentication/login.html")
    return render_template("./Authentication/login.html")


@app.route("/recovery-email", methods=["GET", "POST"])
def recoveryemail():
    email = request.form["email"]
    db.execute(
        "UPDATE users SET recovery=:email WHERE id=:userid",
        {"email": email, "userid": user_id[0]},
    )
    db.commit()
    return redirect(url_for("user_info"))


@app.route("/forgot-password")
def forgotpassword():
    return render_template("./sidebar-09/forgot_password_email.html")


@app.route("/reset-password", methods=["GET", "POST"])
def resetpassword():
    if request.method == "POST":
        global emails
        passs = request.form["password"]
        confirm = request.form["confirmpassword"]
        if (passs is None) or (confirm is None):
            flash("Donot leave blank fields!", "danger")
            return render_template("./sidebar-09/reset_password.html")
        else:
            if passs == confirm:
                secure_password = sha256_crypt.encrypt(str(passs))
                db.execute(
                    "UPDATE users SET password=:passs WHERE recovery=:email",
                    {"passs": secure_password, "email": emails},
                )
                db.commit()
                return redirect(url_for("login"))
            else:
                flash("Password and Confirm Password doesnot match!", "danger")
                return render_template("./sidebar-09/reset_password.html")
    return render_template("./sidebar-09/reset_password.html")


@app.route("/account-found", methods=["POST"])
def accountfound():
    email = request.form["email"]
    e = db.execute(
        "SELECT recovery FROM users WHERE recovery=:recovery", {
            "recovery": email}
    ).fetchone()
    db.commit()
    if e is None:
        flash("Recovery Email not found!", "danger")
        return redirect(url_for("forgotpassword"))
    else:
        token = s.dumps(email)
        msg = Message(
            "Reset Password", sender="datastoolkit@gmail.com", recipients=[email]
        )
        link = url_for("gettoken", token=token, _external=True)
        msg.body = "Your link is {}".format(link)
        mail.send(msg)
        flash("Email has been sent!", "success")
        return redirect(url_for("login"))


@app.route("/get-token/<token>")
def gettoken(token):
    global emails
    emails = s.loads(token, max_age=600)
    return redirect(url_for("resetpassword"))


@app.route("/summary", methods=["GET", "POST"])
def summary():
    global page, summ
    if request.form.get("summary"):
        page = request.form.get("summary")
    if page != "":
        tokens = tokenizer(page)
        sents = sent_tokenizer(page)
        word_counts = count_words(tokens)
        freq_dist = word_freq_distribution(word_counts)
        sent_scores = score_sentences(sents, freq_dist)
        if request.form.get("sentences"):
            number = request.form.get("sentences")
            number = int(number)
            doc = nlargest(number, sent_scores, key=sent_scores.get)
        else:
            doc = nlargest(10, sent_scores, key=sent_scores.get)
        doc = [word for word in doc]
        doc = " ".join(doc)
        doc = "".join([s for s in doc.strip().splitlines(True) if s.strip()])
        summ = doc
        # name = db.execute("SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
        # profile = db.execute("SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
        # cover = db.execute("SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
        return render_template("./sidebar-09/summary.html", summ=summ, page=page)
    else:
        flash("Upload file or enter text!", "danger")
        return redirect(url_for("textsummarization"))


@app.route("/analyse", methods=["GET", "POST"])
def analyse():
    global page, emotions, intensity, positive_sentences, positive, negative, neutral, x, y, negative_sentences, neutral_sentences
    if request.form.get("sentiment"):
        page = request.form.get("analyse")
        if page != "":
            count = 0
            pos_correct = 0
            neg_correct = 0
            neu_correct = 0
            pos_dic = {}
            neg_dic = {}
            neu_dic = {}
            sents = []
            analyzer = SentimentIntensityAnalyzer()
            for sent in page.split("."):
                sent = sent.strip()
                if sent == "" or None:
                    continue
                if len(sent) < 6:
                    continue
                polarity = TextBlob(sent).sentiment.polarity
                if polarity > 0.5:
                    pos_correct += 1
                    count += 1
                    pos_dic[sent] = polarity
                    sents.append(polarity)
                elif polarity < -0.5:
                    neg_dic[sent] = abs(polarity)
                    sents.append(polarity)
                    neg_correct += 1
                    count += 1
                elif polarity == 0:
                    neu_correct += 1
                    sents.append(0)
                    neu_dic[sent] = 0
                    count += 1
                else:
                    vs = analyzer.polarity_scores(sent)
                    if (not vs["neg"] > 0.1) and (vs["pos"] - vs["neg"]) > 0:
                        pos_correct += 1
                        sents.append(vs["pos"])
                        pos_dic[sent] = vs["pos"]
                        count += 1

                    elif (not vs["pos"] > 0.1) and (vs["pos"] - vs["neg"]) < 0:
                        neg_correct += 1
                        sents.append((-1) * (vs["neg"]))
                        neg_dic[sent] = vs["neg"]
                        count += 1
                    else:
                        neu_correct += 1
                        sents.append(0)
                        neu_dic[sent] = 0
                        count += 1

            positive = (pos_correct / (count)) * 100
            positive = "{:.2f}".format(positive)
            negative = (neg_correct / (count)) * 100
            negative = "{:.2f}".format(negative)
            neutral = (neu_correct / count) * 100
            neutral = "{:.2f}".format(neutral)
            pos_list = list(pos_dic.keys())
            neg_list = list(neg_dic.keys())
            neu_list = list(neu_dic.keys())
            ps = len(pos_list)
            ng = len(neg_list)
            nu = len(neu_list)
            positive_sentences = ""
            negative_sentences = ""
            neutral_sentences = ""
            i = 0
            while i < ps:
                positive_sentences += " " + pos_list[i]
                i += 1
            i = 0
            while i < ng:
                negative_sentences += " " + neg_list[i]
                i += 1
            i = 0
            while i < nu:
                neutral_sentences += " " + neu_list[i]
                i += 1
            i = 0
            x = []
            y = []
            while i < (count):
                x.append(i)
                i = i + 1
            i = 0
            while i < len(sents):
                y.append(sents[i])
                i = i + 1
            fig = Figure()
            fig.set_figheight(4)
            fig.set_figwidth(14)
            axis = fig.add_subplot(1, 1, 1)
            axis.set_title("Sentences Sentiment")
            axis.set_xlabel("Sentences")
            axis.set_ylabel("Polarity")
            axis.set_ylim([-1, 1])
            axis.plot(x, y)

            fi = Figure()
            fi.set_figheight(4)
            fi.set_figwidth(14)
            axi = fi.add_subplot(1, 1, 1)
            axi.set_title("Document Sentiment")
            axi.pie(
                [positive, negative, neutral],
                autopct="%1.2f%%",
                colors=["green", "red", "gray"],
                shadow=True,
                labels=["Positive", "Negative", "Neutral"],
            )
            axi.legend(bbox_to_anchor=(1.5, 1.025), loc="upper left")

            pngImage = io.BytesIO()
            ngImage = io.BytesIO()
            FigureCanvas(fig).print_png(pngImage)
            FigureCanvas(fi).print_png(ngImage)
            pngImageB64String = "data:image/png;base64,"
            ngImageB64String = "data:image/png;base64,"
            pngImageB64String += base64.b64encode(
                pngImage.getvalue()).decode("utf8")
            ngImageB64String += base64.b64encode(
                ngImage.getvalue()).decode("utf8")
            return render_template(
                "./sidebar-09/analyse.html",
                image=pngImageB64String,
                image1=ngImageB64String,
                positive=positive,
                negative=negative,
                neutral=neutral,
                ps=positive_sentences,
                nes=negative_sentences,
                nus=neutral_sentences,
                page=page,
            )
        else:
            flash("Upload file or enter text!", "danger")
            return redirect(url_for("sentimentanalysis"))
    if request.form.get("tone"):
        page = request.form.get("analyse")
        if page != "":
            try:
                emotions = []
                intensity = []
                res = ta.tone(page).get_result()
                elen = len(res["document_tone"]["tones"])
                i = 0
                while i < elen:
                    emotions.append(res["document_tone"]
                                    ["tones"][i]["tone_name"])
                    intensity.append(res["document_tone"]["tones"][i]["score"])
                    i += 1
                i = 0
                while i < len(intensity):
                    intensity[i] = intensity[i] * 100
                    i += 1
                fi = Figure()
                fi.set_figheight(4)
                fi.set_figwidth(14)
                axi = fi.add_subplot(1, 1, 1)
                axi.set_title("Document Tone")
                axi.bar(emotions, intensity)
                pngImage = io.BytesIO()
                FigureCanvas(fi).print_png(pngImage)
                pngImageB64String = "data:image/png;base64,"
                pngImageB64String += base64.b64encode(pngImage.getvalue()).decode(
                    "utf8"
                )
                return render_template(
                    "./sidebar-09/tone.html", image=pngImageB64String, page=page
                )
            except:
                flash("File size is too big!", "danger")
                return redirect(url_for("sentimentanalysis"))

        else:
            flash("Upload file or enter text!", "danger")
            return redirect(url_for("sentimentanalysis"))


@app.route("/topics", methods=["GET", "POST"])
def topics():
    global pages_length, topic_pages, topicss
    topicss = []
    page = request.form.get("topic")
    if page != "":
        if pages_length == 1:
            tokens = []
            data = []
            for word in page.split(" "):
                tokens.append(word.strip().lower())
            for token in tokens:
                if token not in stop_words and token not in punctuation:
                    data.append(token)
            data = [re.sub("\S*@\S*\s?", "", sent) for sent in data]
            data = [re.sub("\s+", " ", sent) for sent in data]
            data = [re.sub("'", "", sent) for sent in data]
            result = []
            for token in data:
                if (
                    token not in gensim.parsing.preprocessing.STOPWORDS
                    and len(token) > 3
                ):
                    result.append(token)
            dictionary = corpora.Dictionary([result])
            corpus = [dictionary.doc2bow(text) for text in [result]]
            model = gensim.models.ldamodel.LdaModel(
                corpus=corpus,
                id2word=dictionary,
                num_topics=pages_length,
                random_state=100,
                update_every=1,
                chunksize=100,
                passes=10,
                alpha="auto",
                per_word_topics=True,
            )
            for topic_id in range(model.num_topics):
                topk = model.show_topic(topic_id, 10)
                topk_words = [w for w, _ in topk]

                print("{}: {}".format(topic_id, " ".join(topk_words)))
                pages_length = 1
                topicss.append(topk_words)
                return render_template(
                    "./sidebar-09/model.html", page=page, topic=topicss
                )
        else:
            for page in topic_pages:
                if page != "":
                    tokens = []
                    data = []
                    for word in page.split(" "):
                        tokens.append(word.strip().lower())
                    for token in tokens:
                        if token not in stop_words and token not in punctuation:
                            data.append(token)
                    data = [re.sub("\S*@\S*\s?", "", sent) for sent in data]
                    data = [re.sub("\s+", " ", sent) for sent in data]
                    data = [re.sub("'", "", sent) for sent in data]
                    result = []
                    for token in data:
                        if (
                            token not in gensim.parsing.preprocessing.STOPWORDS
                            and len(token) > 3
                        ):
                            result.append(token)
                    dictionary = corpora.Dictionary([result])
                    corpus = [dictionary.doc2bow(text) for text in [result]]
                    model = gensim.models.ldamodel.LdaModel(
                        corpus=corpus,
                        id2word=dictionary,
                        num_topics=1,
                        random_state=100,
                        update_every=1,
                        chunksize=100,
                        passes=10,
                        alpha="auto",
                        per_word_topics=True,
                    )
                    for topic_id in range(model.num_topics):
                        topk = model.show_topic(topic_id, 10)
                        topk_words = [w for w, _ in topk]

                        print("{}: {}".format(topic_id, " ".join(topk_words)))
                        topicss.append(topk_words)
            topic_pages = []
            pages_length = 1
            return render_template("./sidebar-09/model.html", page=page, topic=topicss)
    else:
        flash("Upload file or enter text!", "danger")
        return redirect(url_for("topicmodeling"))


@app.route("/confirm-summary-save")
def savesummary():
    global page, summ
    return render_template("./sidebar-09/pdf.html", page=page, summ=summ)


@app.route("/confirm-sentiment-save")
def confirmsentiment():
    global page, positive_sentences, negative_sentences, neutral_sentences, x, y, positive, negative, neutral
    if positive_sentences == "":
        positive_sentences = "None"
    if negative_sentences == "":
        negative_sentences = "None"
    if neutral_sentences == "":
        neutral_sentences = "None"
    fig = Figure()
    fig.set_figheight(2)
    fig.set_figwidth(8)
    axis = fig.add_subplot(1, 1, 1)
    axis.set_title("Sentences Sentiment")
    axis.set_xlabel("Sentences")
    axis.set_ylabel("Polarity")
    axis.set_ylim([-1, 1])
    axis.plot(x, y)

    fi = Figure()
    fi.set_figheight(2)
    fi.set_figwidth(8)
    axi = fi.add_subplot(1, 1, 1)
    axi.set_title("Document Sentiment")
    axi.pie(
        [positive, negative, neutral],
        autopct="%1.2f%%",
        colors=["green", "red", "gray"],
        shadow=True,
        labels=["Positive", "Negative", "Neutral"],
    )
    axi.legend(bbox_to_anchor=(1.5, 1.025), loc="upper left")

    pngImage = io.BytesIO()
    ngImage = io.BytesIO()
    FigureCanvas(fig).print_png(pngImage)
    FigureCanvas(fi).print_png(ngImage)
    pngImageB64String = "data:image/png;base64,"
    ngImageB64String = "data:image/png;base64,"
    pngImageB64String += base64.b64encode(pngImage.getvalue()).decode("utf8")
    ngImageB64String += base64.b64encode(ngImage.getvalue()).decode("utf8")
    return render_template(
        "./sidebar-09/sentimentpdf.html",
        page=page,
        ps=positive_sentences,
        nes=negative_sentences,
        nus=neutral_sentences,
        pngImageB64String=pngImageB64String,
        ngImageB64String=ngImageB64String,
        positive=positive,
        negative=negative,
        neutral=neutral
    )


@app.route("/confirm-tone-save")
def confirmtone():
    global emotions, intensity
    fi = Figure()
    fi.set_figheight(2)
    fi.set_figwidth(8)
    axi = fi.add_subplot(1, 1, 1)
    axi.set_title("Document Tone")
    axi.bar(emotions, intensity)
    pngImage = io.BytesIO()
    FigureCanvas(fi).print_png(pngImage)
    pngImageB64String = "data:image/png;base64,"
    pngImageB64String += base64.b64encode(pngImage.getvalue()).decode("utf8")
    global page, summ
    return render_template(
        "./sidebar-09/tonepdf.html", page=page, image=pngImageB64String
    )


@app.route("/confirm-modeling-save")
def savemodeling():
    global page, topicss
    return render_template("./sidebar-09/topicpdf.html", page=page, summ=topicss)


@app.route("/save", methods=["GET", "POST"])
def save():
    name = request.form.get("fname")
    if name != "":
        name = name + ".pdf"
        pdf = weasyprint.HTML(
            "http://127.0.0.1:5000/confirm-summary-save").write_pdf()
        open(
            os.path.join(
                "C:/Users/HP/Documents/GitHub/DataScienceToolkit-60percent/FYP/flask/data/reports",
                name,
            ),
            "wb",
        ).write(pdf)
        db.execute(
            "INSERT INTO reports(user_id, path) VALUES(:user_id, :path)",
            {"user_id": user_id[0], "path": name},
        )
        db.commit()
        return redirect(url_for("reports"))
    else:
        flash("Enter file name to save!", "danger")
        return redirect(url_for("summary"))


@app.route("/save-sentiment", methods=["GET", "POST"])
def savesentiment():
    name = request.form.get("fname")
    if name != "":
        name = name + ".pdf"
        pdf = weasyprint.HTML(
            "http://127.0.0.1:5000/confirm-sentiment-save"
        ).write_pdf()
        open(
            os.path.join(
                "C:/Users/HP/Documents/GitHub/DataScienceToolkit-60percent/FYP/flask/data/reports",
                name,
            ),
            "wb",
        ).write(pdf)
        db.execute(
            "INSERT INTO reports(user_id, path) VALUES(:user_id, :path)",
            {"user_id": user_id[0], "path": name},
        )
        db.commit()
        return redirect(url_for("reports"))
    else:
        flash("Enter file name to save!", "danger")
        return redirect(url_for("reports"))


@app.route("/save-tone", methods=["GET", "POST"])
def savetone():
    name = request.form.get("fname")
    if name != "":
        name = name + ".pdf"
        pdf = weasyprint.HTML(
            "http://127.0.0.1:5000/confirm-tone-save").write_pdf()
        open(
            os.path.join(
                "C:/Users/HP/Documents/GitHub/DataScienceToolkit-60percent/FYP/flask/data/reports",
                name,
            ),
            "wb",
        ).write(pdf)
        db.execute(
            "INSERT INTO reports(user_id, path) VALUES(:user_id, :path)",
            {"user_id": user_id[0], "path": name},
        )
        db.commit()
        return redirect(url_for("reports"))
    else:
        flash("Enter file name to save!", "danger")
        return redirect(url_for("reports"))


@app.route("/save-model", methods=["GET", "POST"])
def savemodel():
    name = request.form.get("fname")
    if name != "":
        name = name + ".pdf"
        pdf = weasyprint.HTML(
            "http://127.0.0.1:5000/confirm-modeling-save").write_pdf()
        open(
            os.path.join(
                "C:/Users/HP/Documents/GitHub/DataScienceToolkit-60percent/FYP/flask/data/reports",
                name,
            ),
            "wb",
        ).write(pdf)
        db.execute(
            "INSERT INTO reports(user_id, path) VALUES(:user_id, :path)",
            {"user_id": user_id[0], "path": name},
        )
        db.commit()
        return redirect(url_for("reports"))
    else:
        flash("Enter file name to save!", "danger")
        return redirect(url_for("reports"))


@app.route("/reports")
def reports():
    paths = db.execute(
        "SELECT * FROM reports WHERE user_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    print(paths)
    return render_template("./sidebar-09/reports.html", paths=paths)


@app.route("/home")
def home():
    return render_template("./sidebar-09/home1.html")


@app.route("/text-summarization")
def textsummarization():
    global page
    # name = db.execute("SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    # profile = db.execute("SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    # cover = db.execute("SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    return render_template("./sidebar-09/text_summarization.html", page=page)


@app.route("/sentiment-analysis")
def sentimentanalysis():
    return render_template("./sidebar-09/sentiment_analysis.html", page=page)


@app.route("/topic-modeling")
def topicmodeling():
    return render_template("./sidebar-09/topic_modeling.html", page=page)


@app.route("/clean-for-textsummarization")
def cleantext():
    global page
    page = ""
    return redirect(url_for("textsummarization"))


@app.route("/clean-for-sentimentanalysis")
def cleansentiment():
    global page
    page = ""
    return redirect(url_for("sentimentanalysis"))


@app.route("/clean-for-topicmodeling")
def cleantopic():
    global page
    page = ""
    return redirect(url_for("topicmodeling"))


@app.route("/data-cleaning", methods=["POST", "GET"])
def datacleaning():
    global page
    option = request.form.get("cleaning")
    if option == "lower":
        page = page.lower()
    elif option == "upper":
        page = page.upper()
    elif option == "proper":
        page = page.title()
    elif option == "punctuation":
        pun = """!()-[]{};:'"\,<>./?@#$%^&*_~"""
        no_punct = ""
        for char in page:
            if char not in pun:
                no_punct = no_punct + char
        page = no_punct
    else:
        page = "".join([s for s in page.strip().splitlines(True) if s.strip()])
    if request.form.get("textsummarization"):
        return redirect(url_for("textsummarization"))
    elif request.form.get("sentimentanalysis"):
        return redirect(url_for("sentimentanalysis"))
    elif request.form.get("topicmodeling"):
        return redirect(url_for("topicmodeling"))


@app.route("/upload-file", methods=["POST"])
def upload():
    global page, topic_pages
    global f_name, pages_length
    page = ""
    pages_length = 1
    topic_pages = []
    filess = request.files.getlist("fileToUpload")
    pages_length = len(filess)
    for files in filess:
        if files.filename == "":
            flash("Select file!", "danger")
            break
    for files in filess:
        filename = files.filename
        farray = filename.split(".")
        length = len(farray)
        ftype = length - 1
        filetype = farray[ftype]
        filetype = filetype.lower()
        if (
            filetype == "png"
            or filetype == "jpg"
            or filetype == "jpeg"
            or filetype == "bmp"
            or filetype == "gif"
            or filetype == "tif"
        ):
            files.save(os.path.join("/FYP/flask/data/files", filename))
            f_name = filename
            pytesseract.pytesseract.tesseract_cmd = (
                r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            )
            img = Image.open(files)
            text = pytesseract.image_to_string(img)
            text = "".join(
                [s for s in text.strip().splitlines(True) if s.strip()])
            page += text
            topic_pages.append(text)

        elif filetype == "txt":
            files.save(os.path.join("/FYP/flask/data/files", filename))
            f_name = filename
            f = open(os.path.join("/FYP/flask/data/files", filename), "r")
            text = f.read()
            page += text
            topic_pages.append(text)

        elif filetype == "pdf":
            try:
                files.save(os.path.join("/FYP/flask/data/files", filename))
                f_name = filename
                pdfReader = PyPDF2.PdfFileReader(files)
                pages = pdfReader.numPages
                i = 0
                text = ""
                while i < pages:
                    pageObject = pdfReader.getPage(i)
                    text += pageObject.extractText()
                    i = i + 1
                text1 = "".join(
                    [s for s in text.strip().splitlines(True) if s.strip()])
                page += text1
                topic_pages.append(text1)
            except:
                continue

        elif filetype == "docx" or filetype == "doc":
            files.save(os.path.join("/FYP/flask/data/files", filename))
            f_name = filename
            doc = docx.Document(files)
            paragraphs = len(doc.paragraphs)
            i = 0
            text = ""
            while i < paragraphs:
                text += doc.paragraphs[i].text
                i = i + 1
                text += "\n"
            page += "".join([s for s in text.strip().splitlines(True) if s.strip()])
            topic_pages.append(
                "".join([s for s in text.strip().splitlines(True) if s.strip()])
            )
        elif filetype == "epub":
            files.save(os.path.join("/FYP/flask/data/files", filename))
            f_name = filename

            def epub2thtml(epub_path):
                book = epub.read_epub(epub_path)
                chapters = []
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        chapters.append(item.get_content())
                return chapters

            def chap2text(chap):
                output = ""
                soup = BeautifulSoup(chap, "html.parser")
                text = soup.find_all(text=True)
                blacklist = [
                    "[document]",
                    "noscript",
                    "header",
                    "html",
                    "meta",
                    "head",
                    "input",
                    "script",
                ]
                for t in text:
                    if t.parent.name not in blacklist:
                        output += "{} ".format(t)
                return output

            def thtml2ttext(thtml):
                Output = []
                for html in thtml:
                    text = chap2text(html)
                    Output.append(text)
                return Output

            def epub2text(epub_path):
                chapters = epub2thtml(epub_path)
                ttext = thtml2ttext(chapters)
                return ttext

            text = epub2text(files)
            length = len(text)
            i = 0
            text1 = ""
            while i < length:
                text1 += text[i]
                i = i + 1
            text1 = "".join(
                [s for s in text1.strip().splitlines(True) if s.strip()])
            page += text1
            topic_pages.append(text1)

        elif filetype == "mobi":
            book = mobi_uncompress(files)
            book.parse()
            text = ""
            for record in book:
                text += record
            page += text
            topic_pages.append(text)
    if request.form["submit"] == "Submit file for text-summarization":
        return redirect(url_for("textsummarization"))
    elif request.form["submit"] == "Submit file for sentiment-analysis":
        return redirect(url_for("sentimentanalysis"))
    elif request.form["submit"] == "Submit files for topic-modeling":
        return redirect(url_for("topicmodeling"))
    elif request.form["submit"] == "save":
        if files.filename != "":
            filename = secure_filename(files.filename)
            name = f_name.split(".")
            name = name[0] + ".txt"
            textfile = open(os.path.join("/FYP/flask/data/txt", name), "w")
            textfile.write(page)
            textfile.close()
            db.execute(
                "INSERT INTO files(user_id, path, type) VALUES(:user_id, :path, :type)",
                {"user_id": user_id[0], "path": f_name, "type": filetype},
            )
            db.commit()
            flash("File saved!", "success")
            return redirect(url_for("repository"))
        else:
            return redirect(url_for("repository"))
    return("File not uploaded")


@app.route("/web-scraping", methods=["POST"])
def web():
    global page
    url = request.form["url"]
    if url != "":
        html = urlopen(url).read()
        soup = BeautifulSoup(html)
        for script in soup(["script", "style"]):
            script.extract()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip()
                  for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        page = text

    else:
        flash("Enter url!", "danger")
    if request.form.get("textsummarization"):
        return redirect(url_for("textsummarization"))
    elif request.form.get("sentimentanalysis"):
        return redirect(url_for("sentimentanalysis"))
    elif request.form.get("topicmodeling"):
        return redirect(url_for("topicmodeling"))


@app.route("/show-repository-for-text-summarization")
def showrepositorytext():
    name = db.execute(
        "SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    profile = db.execute(
        "SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    cover = db.execute(
        "SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    path = db.execute(
        "SELECT path FROM files WHERE user_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    path1 = db.execute(
        "SELECT path FROM links WHERE second_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    return render_template(
        "./sidebar-09/show_repository_for_textsummarization.html",
        path1=path1,
        n=name[0],
        p=profile[0],
        c=cover[0],
        path=path,
    )


@app.route("/show-repository-for-sentiment-analysis")
def showrepositorysentiment():
    name = db.execute(
        "SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    profile = db.execute(
        "SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    cover = db.execute(
        "SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    path = db.execute(
        "SELECT path FROM files WHERE user_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    path1 = db.execute(
        "SELECT path FROM links WHERE second_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    return render_template(
        "./sidebar-09/show_repository_for_sentimentanalysis.html",
        path1=path1,
        n=name[0],
        p=profile[0],
        c=cover[0],
        path=path,
    )


@app.route("/show-repository-for-topic-modeling")
def showrepositorytopic():
    name = db.execute(
        "SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    profile = db.execute(
        "SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    cover = db.execute(
        "SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    path = db.execute(
        "SELECT path FROM files WHERE user_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    path1 = db.execute(
        "SELECT path FROM links WHERE second_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    return render_template(
        "./sidebar-09/show_repository_for_topic_modeling.html",
        path1=path1,
        n=name[0],
        p=profile[0],
        c=cover[0],
        path=path,
    )


@app.route("/delete-file/<string:path>", methods=["GET", "POST"])
def deletefile(path):
    db.execute("DELETE FROM files WHERE path=:path", {"path": path})
    db.commit()
    return redirect(url_for("repository"))


@app.route("/delete-files/<string:path>", methods=["GET", "POST"])
def deletefiles(path):
    db.execute("DELETE FROM reports WHERE path=:path", {"path": path})
    db.commit()
    return redirect(url_for("reports"))


@app.route(
    "/upload-from-repository-textsummarization/<string:path>", methods=["GET", "POST"]
)
def uploadrepositorytext(path):
    global page
    name = path.split(".")
    name = name[0] + ".txt"
    f = open(os.path.join("/FYP/flask/data/txt", name), "r")
    page = f.read()
    return redirect(url_for("textsummarization"))


@app.route(
    "/upload-from-repository-sentiment-analysis/<string:path>", methods=["GET", "POST"]
)
def uploadrepositorysentiment(path):
    global page
    name = path.split(".")
    name = name[0] + ".txt"
    f = open(os.path.join("/FYP/flask/data/txt", name), "r")
    page = f.read()
    return redirect(url_for("sentimentanalysis"))


@app.route(
    "/upload-from-repository-topicmodeling/<string:path>", methods=["GET", "POST"]
)
def uploadrepositorytopic(path):
    global page
    name = path.split(".")
    name = name[0] + ".txt"
    f = open(os.path.join("/FYP/flask/data/txt", name), "r")
    page = f.read()
    return redirect(url_for("topicmodeling"))


@app.route("/repository")
def repository():
    name = db.execute(
        "SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    profile = db.execute(
        "SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    cover = db.execute(
        "SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    path = db.execute(
        "SELECT path FROM files WHERE user_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    path1 = db.execute(
        "SELECT first_id FROM links WHERE second_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    i = 1
    names = []
    if not path1:
        return render_template(
            "./sidebar-09/repository.html",
            names=names,
            path1=path1,
            p=profile[0],
            n=name[0],
            c=cover[0],
            path=path,
        )
    else:
        j = path1[0][0]
        l = [23]
        l[0] = j
        while i < len(path1):
            if j != path1[i][0]:
                l.append(path1[i][0])
                j = path1[i][0]
            i = i + 1
        i = 0
        while i < len(l):
            names.append(
                db.execute(
                    "SELECT name,id FROM users WHERE id =:user_id", {
                        "user_id": l[i]}
                ).fetchone()
            )
            db.commit()
            i = i + 1
    return render_template(
        "./sidebar-09/repository.html",
        names=names,
        path1=path1,
        p=profile[0],
        n=name[0],
        c=cover[0],
        path=path,
    )


@app.route("/repository/<string:id>", methods=["GET", "POST"])
def uploadrepositoryname(id):
    # name = db.execute("SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    # profile = db.execute("SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    # cover = db.execute("SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}).fetchone()
    path = db.execute(
        "SELECT path FROM links WHERE first_id =:userid && second_id =:user_id",
        {"userid": id, "user_id": user_id[0]},
    ).fetchall()
    return render_template("./sidebar-09/repository_name.html", path=path)


@app.route("/view-file/<string:path>")
def view(path):
    return send_from_directory(app.config["UPLOAD_FOLDER"], path)


@app.route("/view-files/<string:path>")
def views(path):
    return send_from_directory(
        "C:/Users/HP/Documents/GitHub/DataScienceToolkit-60percent/FYP/flask/data/reports",
        path,
    )


@app.route("/download-files/<string:path>")
def download(path):
    return send_file(
        os.path.join(
            "C:/Users/HP/Documents/GitHub/DataScienceToolkit-60percent/FYP/flask/data/reports",
            path,
        ),
        as_attachment=True,
    )


@app.route("/profile")
def user_info():
    email = db.execute(
        "SELECT recovery FROM users WHERE id =:user_id", {
            "user_id": user_id[0]}
    ).fetchone()
    name = db.execute(
        "SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    profile = db.execute(
        "SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    cover = db.execute(
        "SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    return render_template(
        "./sidebar-09/profile.html", r=email[0], p=profile[0], n=name[0], c=cover[0]
    )


@app.route("/profile-picture", methods=["POST"])
def profile():
    files = request.files["profile"]
    filename = files.filename
    filename = secure_filename(files.filename)
    farray = filename.split(".")
    length = len(farray)
    ftype = length - 1
    filetype = farray[ftype]
    filetype = filetype.lower()
    profile = filename
    if filetype == "png" or filetype == "jpg" or filetype == "jpeg":
        output_size = (400, 400)
        i = Image.open(files)
        i.thumbnail(output_size)
        i.save(os.path.join("/FYP/flask/static/sidebar/images/", filename))
        user1 = user_id[0]
        db.execute(
            "UPDATE users SET profile=:filename WHERE id=:user1",
            {"filename": filename, "user1": user1},
        )
        db.commit()
        return redirect(url_for("user_info"))
    else:
        flash("Only png, jpg and jpeg files are accepted!", "danger")
        return redirect(url_for("user_info"))


@app.route("/cover", methods=["POST"])
def cover():
    files = request.files["cover"]
    filename = files.filename
    filename = secure_filename(files.filename)
    farray = filename.split(".")
    length = len(farray)
    ftype = length - 1
    filetype = farray[ftype]
    filetype = filetype.lower()
    profile = filename
    if filetype == "png" or filetype == "jpg" or filetype == "jpeg":
        output_size = (400, 400)
        i = Image.open(files)
        i.thumbnail(output_size)
        i.save(os.path.join("/FYP/flask/static/sidebar/images/", filename))
        user1 = user_id[0]
        db.execute(
            "UPDATE users SET cover=:filename WHERE id=:user1",
            {"filename": filename, "user1": user1},
        )
        db.commit()
        return redirect(url_for("user_info"))
    else:
        flash("Only png, jpg and jpeg files are accepted!", "danger")
        return redirect(url_for("user_info"))


@app.route("/users-search", methods=["GET", "POST"])
def users_search():
    user_name = request.form.get("search")
    user_name = user_name.split()
    words = len(user_name)
    i = words
    searched = []
    while words > 0:
        if not searched:
            if words == 1:
                searched = db.execute(
                    "SELECT * FROM users WHERE name LIKE '%" +
                    user_name[0] + "%'  "
                ).fetchall()
            elif words == 2:
                searched = db.execute(
                    "SELECT * FROM users WHERE name LIKE '%"
                    + user_name[0]
                    + "%' AND name LIKE '%"
                    + user_name[1]
                    + "%' "
                ).fetchall()
            elif words == 3:
                searched = db.execute(
                    "SELECT * FROM users WHERE name LIKE '%"
                    + user_name[0]
                    + "%' AND name LIKE '%"
                    + user_name[1]
                    + "%' AND name LIKE '%"
                    + user_name[2]
                    + "%' "
                ).fetchall()
            elif words == 4:
                searched = db.execute(
                    "SELECT * FROM users WHERE name LIKE '%"
                    + user_name[0]
                    + "%' AND name LIKE '%"
                    + user_name[1]
                    + "%' AND name LIKE '%"
                    + user_name[2]
                    + "%' AND name LIKE '%"
                    + user_name[3]
                    + "%' "
                ).fetchall()
            elif words == 5:
                searched = db.execute(
                    "SELECT * FROM users WHERE name LIKE '%"
                    + user_name[0]
                    + "%' AND name LIKE '%"
                    + user_name[1]
                    + "%' AND name LIKE '%"
                    + user_name[2]
                    + "%' AND name LIKE '%"
                    + user_name[3]
                    + "%' AND name LIKE '%"
                    + user_name[4]
                    + "%'  "
                ).fetchall()
            else:
                searched = db.execute(
                    "SELECT * FROM users WHERE name LIKE '%"
                    + user_name[0]
                    + "%' AND name LIKE '%"
                    + user_name[1]
                    + "%' AND name LIKE '%"
                    + user_name[2]
                    + "%' AND name LIKE '%"
                    + user_name[3]
                    + "%' AND name LIKE '%"
                    + user_name[4]
                    + "%' AND name LIKE '%"
                    + user_name[5]
                    + "%' "
                ).fetchall()
        words = words - 1
    if not searched:
        if i == 1:
            searched = db.execute(
                "SELECT * FROM users WHERE name LIKE '%" +
                user_name[0] + "%'  "
            ).fetchall()
        elif i == 2:
            searched = db.execute(
                "SELECT * FROM users WHERE name LIKE '%"
                + user_name[0]
                + "%' OR name LIKE '%"
                + user_name[1]
                + "%' "
            ).fetchall()
        elif i == 3:
            searched = db.execute(
                "SELECT * FROM users WHERE name LIKE '%"
                + user_name[0]
                + "%' OR name LIKE '%"
                + user_name[1]
                + "%' OR name LIKE '%"
                + user_name[2]
                + "%' "
            ).fetchall()
        elif i == 4:
            searched = db.execute(
                "SELECT * FROM users WHERE name LIKE '%"
                + user_name[0]
                + "%' OR name LIKE '%"
                + user_name[1]
                + "%' OR name LIKE '%"
                + user_name[2]
                + "%' OR name LIKE '%"
                + user_name[3]
                + "%' "
            ).fetchall()
        elif i == 5:
            searched = db.execute(
                "SELECT * FROM users WHERE name LIKE '%"
                + user_name[0]
                + "%' OR name LIKE '%"
                + user_name[1]
                + "%' OR name LIKE '%"
                + user_name[2]
                + "%' OR name LIKE '%"
                + user_name[3]
                + "%' OR name LIKE '%"
                + user_name[4]
                + "%'  "
            ).fetchall()
        else:
            searched = db.execute(
                "SELECT * FROM users WHERE name LIKE '%"
                + user_name[0]
                + "%' OR name LIKE '%"
                + user_name[1]
                + "%' OR name LIKE '%"
                + user_name[2]
                + "%' OR name LIKE '%"
                + user_name[3]
                + "%' OR name LIKE '%"
                + user_name[4]
                + "%' OR name LIKE '%"
                + user_name[5]
                + "%' "
            ).fetchall()

    name = db.execute(
        "SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    profile = db.execute(
        "SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    cover = db.execute(
        "SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    return render_template(
        "./sidebar-09/searched_users.html",
        n=name[0],
        p=profile[0],
        c=cover[0],
        searched=searched,
    )


@app.route("/<string:user12>")
def user_12(user12):
    name = db.execute(
        "SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    profile = db.execute(
        "SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    cover = db.execute(
        "SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    user = db.execute("SELECT * FROM users WHERE id=:id",
                      {"id": user12}).fetchone()
    search = db.execute(
        "SELECT * FROM follow WHERE first_id =:user12 AND second_id =:user_id",
        {"user12": user12, "user_id": user_id[0]},
    ).fetchone()
    if search is None:
        path = False
    else:
        path = db.execute(
            "SELECT path FROM files WHERE user_id =:user_id", {
                "user_id": user[0]}
        ).fetchall()
    return render_template(
        "./sidebar-09/searched_profile.html",
        n=name[0],
        p=profile[0],
        c=cover[0],
        u=user,
        path=path,
    )


@app.route("/save-to-repository/<string:userid1>")
def savetorepository(userid1):
    db.execute(
        "DELETE FROM links WHERE first_id=:first_id and second_id=:second_id ",
        {"first_id": userid1, "second_id": user_id[0]},
    )
    db.commit()
    path = db.execute(
        "SELECT path FROM files WHERE user_id =:user_id", {"user_id": userid1}
    ).fetchall()
    length = len(path)
    i = 0
    while i < length:
        db.execute(
            "INSERT INTO links(first_id, path, second_id) VALUES(:first_id, :path, :second_id)",
            {"first_id": userid1, "path": path[i][0], "second_id": user_id[0]},
        )
        db.commit()
        i = i + 1
    return redirect(url_for("repository"))


@app.route("/follow/<string:userid12>")
def follow(userid12):
    db.execute(
        "DELETE FROM notification WHERE first_id=:first_id and second_id=:second_id ",
        {"first_id": user_id[0], "second_id": userid12},
    )
    db.commit()
    db.execute(
        "INSERT INTO follow(first_id, second_id) VALUES(:first_id, :second_id)",
        {"first_id": user_id[0], "second_id": userid12},
    )
    db.commit()
    return redirect(url_for("receivenotification"))


@app.route("/send-notification/<string:userid13>")
def sendnotification(userid13):
    name = db.execute(
        "SELECT * FROM follow WHERE first_id =:userid13 AND second_id =:user_id",
        {"userid13": userid13, "user_id": user_id[0]},
    ).fetchone()
    if name is None:
        db.execute(
            "DELETE FROM notification WHERE first_id=:first_id and second_id=:second_id ",
            {"first_id": userid13, "second_id": user_id[0]},
        )
        db.commit()
        db.execute(
            "INSERT INTO notification(first_id, second_id) VALUES(:first_id, :second_id)",
            {"first_id": userid13, "second_id": user_id[0]},
        )
        db.commit()
        return redirect(url_for("feed"))
    else:
        print("USER ALREDY FOLLOWED")
        return redirect(url_for("feed"))


@app.route("/receive-notification")
def receivenotification():
    usersid = db.execute(
        "SELECT second_id FROM notification WHERE first_id =:userid",
        {"userid": user_id[0]},
    ).fetchall()
    length = len(usersid)
    i = 0
    users = []
    while i < length:
        users.append(
            db.execute(
                "SELECT * FROM users WHERE id =:userid", {
                    "userid": usersid[i][0]}
            ).fetchone()
        )
        i = i + 1
    return render_template("./sidebar-09/notifications.html", users=users)


@app.route("/feed")
def feed():
    name = db.execute(
        "SELECT name FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    profile = db.execute(
        "SELECT profile FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    cover = db.execute(
        "SELECT cover FROM users WHERE id =:user_id", {"user_id": user_id[0]}
    ).fetchone()
    ids = db.execute(
        "SELECT first_id FROM follow WHERE second_id =:user_id", {
            "user_id": user_id[0]}
    ).fetchall()
    i = 0
    path = []
    followed = []
    while i < len(ids):
        followed.append(
            [
                (
                    db.execute(
                        "SELECT * FROM users WHERE id =:user_id", {
                            "user_id": ids[i][0]}
                    ).fetchone()
                ),
                i,
            ]
        )
        path.append(
            db.execute(
                "SELECT path FROM files WHERE user_id =:user_id", {
                    "user_id": ids[i][0]}
            ).fetchall()
        )
        db.commit()
        i = i + 1
    return render_template(
        "./sidebar-09/feed.html",
        path=path,
        followed=followed,
        n=name[0],
        p=profile[0],
        c=cover[0],
    )


def tokenizer(s):
    tokens = []
    for word in s.split(" "):
        tokens.append(word.strip().lower())
    return tokens


def sent_tokenizer(s):
    sents = []
    for sent in s.split("."):
        sents.append(sent.strip())
    return sents


def count_words(tokens):
    word_counts = {}
    for token in tokens:
        if token not in stop_words and token not in punctuation:
            if token not in word_counts.keys():
                word_counts[token] = 1
            else:
                word_counts[token] += 1
    return word_counts


def word_freq_distribution(word_counts):
    freq_dist = {}
    max_freq = max(word_counts.values())
    for word in word_counts.keys():
        freq_dist[word] = word_counts[word] / max_freq
    return freq_dist


def score_sentences(sents, freq_dist, max_len=40):
    sent_scores = {}
    for sent in sents:
        words = sent.split(" ")
        for word in words:
            if word.lower() in freq_dist.keys():
                if len(words) < max_len:
                    if sent not in sent_scores.keys():
                        sent_scores[sent] = freq_dist[word.lower()]
                    else:
                        sent_scores[sent] += freq_dist[word.lower()]
    return sent_scores


if __name__ == "__main__":
    app.secret_key = ""
    app.run(debug=True)
